"""Generate HANDOFF.docx for Advert Radical project."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x26, 0x34, 0x31)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x26, 0x34, 0x31)
    hs.font.name = 'Georgia'

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x26, 0x34, 0x31)
    return h

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    doc.add_paragraph()

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x3A, 0x4A, 0x46)
    p.paragraph_format.left_indent = Cm(1)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ADVERT RADICAL')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x26, 0x34, 0x31)
run.font.name = 'Georgia'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Project Handoff Document')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xC4, 0xA8, 0x6D)
run.font.name = 'Georgia'
run.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Version 6  |  April 2026  |  Prabu D, Chennai, India')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x63, 0x5F, 0x5A)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Production: ar-saa-s-kbzn.vercel.app\nGitHub: github.com/git-prabu/AR-SaaS')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x63, 0x5F, 0x5A)

doc.add_page_break()

# ============================================================
# 1. WHAT IS ADVERT RADICAL?
# ============================================================
add_heading('1. What Is Advert Radical?', 1)
add_para('Advert Radical is an AR + AI SaaS platform that gives Indian restaurants a QR-code-based digital menu with:')
add_bullet('WebAR 3D dish previews — customers scan a table QR code, see a live 3D model of each dish on their phone')
add_bullet('AI-powered upselling — Claude (Anthropic) suggests add-ons & combos based on what customers are ordering')
add_bullet('Real-time waiter calls — customers tap to call a waiter; kitchen/staff get instant audio + push notifications')
add_bullet('Full admin dashboard — analytics, menu management, order tracking, combo builder, offers, QR codes, payments')
add_bullet('Super admin panel — platform-level management of all restaurants, plans, AR model requests')

add_para('Business Model', bold=True, size=13)
add_para('B2B SaaS for restaurant chains with three monthly subscription plans:')
add_table(
    ['Plan', 'Price', 'Max Items', 'Storage'],
    [
        ['Starter', '₹999/mo', '25', '500 MB'],
        ['Growth', '₹2,499/mo', '75', '2 GB'],
        ['Pro', '₹4,999/mo', 'Unlimited', '10 GB'],
    ]
)

add_para('Demo Routes', bold=True, size=13)
add_bullet('Customer menu: /restaurant/spot')
add_bullet('Admin dashboard: /admin/login')
add_bullet('Super admin: /superadmin/login')
add_bullet('Pitch deck: /pitch (password: RADICAL25)')

# ============================================================
# 2. TECH STACK
# ============================================================
add_heading('2. Tech Stack', 1)
add_table(
    ['Layer', 'Technology'],
    [
        ['Framework', 'Next.js 16.1.6 (Pages Router ONLY — never App Router)'],
        ['Frontend', 'React 18, CSS-in-JS (inline styles), Recharts, GSAP, react-hot-toast'],
        ['Auth', 'Firebase Auth (dual instances: adminAuth + superAdminAuth)'],
        ['Database', 'Firestore (client SDK via lib/firebase.js)'],
        ['Storage', 'Firebase Storage (images + GLB 3D models)'],
        ['AR', '<model-viewer> web component via public/ar-viewer.html iframe'],
        ['AI', 'Claude API (Anthropic) for menu upselling'],
        ['Payments', 'Razorpay (API routes for order creation + verification)'],
        ['3D Generation', 'Meshy AI Image-to-3D API (generates GLB from dish photos)'],
        ['Drag & Drop', '@dnd-kit (menu item reordering)'],
        ['Export', 'xlsx (CSV/Excel export from analytics)'],
        ['Hosting', 'Vercel (auto-deploy from main branch)'],
    ]
)

# ============================================================
# 3. PROJECT STRUCTURE
# ============================================================
add_heading('3. Project Structure', 1)

add_para('Pages', bold=True, size=13)
structure_items = [
    ('pages/index.js', 'Landing page / homepage'),
    ('pages/pitch.js', 'Investor pitch deck (pw: RADICAL25)'),
    ('pages/seed.js', 'Database seeder (dev utility)'),
    ('pages/_app.js', 'App wrapper — auth providers, error boundary'),
    ('pages/_document.js', 'HTML document head'),
    ('pages/restaurant/[subdomain]/index.js', 'Customer menu page (967 lines) — AR viewer, ratings, combos, orders, waiter calls'),
    ('pages/admin/login.js', 'Restaurant admin login'),
    ('pages/admin/index.js', 'Admin dashboard home'),
    ('pages/admin/analytics.js', 'Analytics dashboard (967 lines) — visits, AR, items, orders, revenue, waiter metrics'),
    ('pages/admin/orders.js', 'Order management (pending → preparing → ready → served)'),
    ('pages/admin/items.js', 'Menu items CRUD (~566 lines)'),
    ('pages/admin/combos.js', 'Smart combo builder'),
    ('pages/admin/offers.js', 'Time-limited offers management'),
    ('pages/admin/payments.js', 'Payment/transaction history'),
    ('pages/admin/requests.js', 'AR model upload requests'),
    ('pages/admin/notifications.js', 'Staff push notifications'),
    ('pages/admin/qrcode.js', 'QR code generator for table sessions'),
    ('pages/admin/subscription.js', 'Plan management & Razorpay integration'),
    ('pages/superadmin/login.js', 'Super admin login'),
    ('pages/superadmin/index.js', 'Platform overview dashboard'),
    ('pages/superadmin/restaurants.js', 'Manage all restaurants'),
    ('pages/superadmin/requests.js', 'Approve/reject AR model requests'),
    ('pages/superadmin/plans.js', 'Subscription plan management'),
    ('pages/superadmin/restaurant/[id].js', 'Individual restaurant detail page'),
]
for path, desc in structure_items:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(f'  —  {desc}')

add_para('')
add_para('API Routes', bold=True, size=13)
api_items = [
    ('pages/api/generate-model.js', 'Meshy AI Image-to-3D API — generates GLB models from dish photos'),
    ('pages/api/payments/create-order.js', 'Create Razorpay payment orders for plan upgrades'),
    ('pages/api/payments/verify.js', 'Verify Razorpay payment signature & update subscription'),
]
for path, desc in api_items:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(f'  —  {desc}')

add_para('')
add_para('Components', bold=True, size=13)
add_bullet('components/layout/AdminLayout.jsx — Admin sidebar with real-time listeners, sound notifications, responsive nav')
add_bullet('components/layout/SuperAdminLayout.jsx — Super admin sidebar (dark theme)')
add_bullet('components/ARViewer.jsx — model-viewer iframe wrapper for AR display')

add_para('')
add_para('Libraries', bold=True, size=13)
lib_items = [
    ('lib/firebase.js', 'Firebase client init — dual apps (adminApp + superAdminApp), exports db, auth, storage'),
    ('lib/firebaseAdmin.js', 'Firebase Admin SDK (server-side for API routes)'),
    ('lib/db.js', 'Client Firestore helpers — 40+ exported functions (514 lines)'),
    ('lib/saDb.js', 'Super admin Firestore helpers (uses superAdminDb)'),
    ('lib/storage.js', 'Firebase Storage upload/delete (admin)'),
    ('lib/saStorage.js', 'Firebase Storage upload/delete (superadmin)'),
    ('lib/utils.js', 'Design tokens (T object) + utility functions'),
]
for path, desc in lib_items:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(f'  —  {desc}')

add_para('')
add_para('Hooks & Contexts', bold=True, size=13)
add_bullet('hooks/useAuth.js — Dual auth hook: useAuth/useAdminAuth for admins, useSuperAdminAuth for superadmins')
add_bullet('contexts/AdminDataContext.js — Shared real-time orders & waiter calls (prevents duplicate Firestore listeners)')

add_para('')
add_para('Public Assets', bold=True, size=13)
add_bullet('public/ar-viewer.html — Standalone model-viewer HTML for AR iframe')
add_bullet('public/notification.mp3 — Bell tone audio for order/call alerts')
add_bullet('public/ar-experience.png — AR experience marketing image')

# ============================================================
# 4. ARCHITECTURE
# ============================================================
add_heading('4. Architecture', 1)

add_para('4.1 Dual Firebase Authentication', bold=True, size=13)
add_para('Two separate Firebase app instances run in the browser to prevent auth collisions:')
add_bullet('adminApp — used by restaurant admins AND customers')
add_bullet('superAdminApp — used exclusively by platform super admins')
add_para('Each has its own auth object with independent sign-in sessions. lib/firebase.js exports: adminApp, adminAuth, superAdminAuth, db, superAdminDb, storage, superAdminStorage.')

add_para('')
add_para('4.2 Subdomain Multi-Tenancy', bold=True, size=13)
add_para('middleware.js intercepts requests at the Vercel edge:')
add_bullet('Extracts subdomain from hostname (e.g., "spot" from spot.advertradical.com)')
add_bullet('Rewrites to /restaurant/[subdomain] — the customer menu page')
add_bullet('Local dev uses ?sub=spotname query param')
add_bullet('Reserved subdomains (www, superadmin, api) are skipped')

add_para('')
add_para('4.3 Real-Time Data Flow', bold=True, size=13)
add_para('AdminLayout.jsx sets up Firestore onSnapshot listeners for orders and waiterCalls collections. These are shared via AdminDataContext to all admin pages, preventing duplicate subscriptions. Admin pages consume data via useAdminOrders() and useAdminWaiterCalls() hooks.')

add_para('')
add_para('4.4 AR Model Pipeline', bold=True, size=13)
add_bullet('Restaurant uploads dish photo')
add_bullet('Admin submits AR request (pages/admin/requests.js)')
add_bullet('Super admin approves (pages/superadmin/requests.js)')
add_bullet('Meshy AI generates GLB model (pages/api/generate-model.js)')
add_bullet('Model stored in Firebase Storage')
add_bullet('Customer sees 3D model in menu (public/ar-viewer.html iframe)')

add_para('')
add_para('4.5 Payment Flow', bold=True, size=13)
add_bullet('Customer places order → createOrder() in Firestore')
add_bullet('Admin creates Razorpay order → /api/payments/create-order')
add_bullet('Razorpay checkout opens → customer pays')
add_bullet('Verify signature → /api/payments/verify')
add_bullet('Update restaurant subscription in Firestore')

# ============================================================
# 5. FIRESTORE DATA MODEL
# ============================================================
add_heading('5. Firestore Data Model', 1)

add_para('restaurants/{restaurantId}', bold=True, size=12)
add_para('Top-level fields: name, subdomain, email, phone, address, logo, plan, maxItems, maxStorageMB, subscriptionStart, subscriptionEnd, paymentStatus, isActive, createdAt')

add_para('')
add_para('Sub-collections:', bold=True, size=12)

collections = [
    ('menuItems/{itemId}', 'name, description, category, price, imageURL, modelURL, arReady, isFeatured, isActive, sortOrder, ingredients[], calories, protein, carbs, fats, prepTime, spiceLevel, isVeg, badge, views, arViews, ratingSum, ratingCount, ratingAvg, createdAt, updatedAt'),
    ('orders/{orderId}', 'items [{name, price, qty}], total, tableNumber, status (pending|preparing|ready|served), paymentStatus (paid|unpaid), paymentUpdatedAt, createdAt, updatedAt'),
    ('analytics/{YYYY-MM-DD}', 'totalVisits, uniqueVisitors, repeatVisitors, sessions[], date'),
    ('waiterCalls/{callId}', 'tableNumber, status (pending|resolved), createdAt, resolvedAt'),
    ('offers/{offerId}', 'title, description, discount, startDate, endDate, applicableItems[], createdAt, updatedAt'),
    ('combos/{comboId}', 'name, description, price, itemIds[], isActive, createdAt'),
    ('tableSessions/{tableNumber}', 'sid (session ID — rotates on QR activation), isActive, expiresAt, createdAt'),
    ('requests/{requestId}', 'name, description, category, price, imageURL, nutritionalData {calories, protein, carbs, fats}, prepTime, spiceLevel, isVeg, badge, ingredients[], status (pending|approved|rejected), modelURL, arReady, createdAt, reviewedAt'),
]

for name, fields in collections:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    doc.add_paragraph(fields)

add_para('')
add_para('Top-level collections:', bold=True, size=12)
add_bullet('users/{uid} — email, role (restaurant|superadmin), restaurantId, restaurantName, createdAt')
add_bullet('plans/{planId} — name, price, maxItems, maxStorageMB, features[]')

# ============================================================
# 6. DESIGN SYSTEM
# ============================================================
add_heading('6. Design System', 1)

add_para('6.1 Cinematic Color Palette (LOCKED)', bold=True, size=13)
add_table(
    ['Token', 'Hex', 'Usage'],
    [
        ['Deep Forest', '#263431', 'Primary accent, sidebar bg, dark sections'],
        ['Antique Gold', '#C4A86D', 'Highlights, badges, gold accents'],
        ['Soft Cream', '#EAE7E3', 'Page backgrounds, muted surfaces'],
        ['Stone Grey', '#635F5A', 'Secondary text, subtle labels'],
        ['White', '#FFFFFF', 'Card backgrounds'],
        ['Charcoal', '#3A4A46', 'Body text on light bg'],
        ['Sand', '#D6CCBA', 'Borders, dividers'],
        ['Success Green', '#4A7A5E', 'Positive metrics'],
        ['Danger Red', '#8A4A42', 'Alerts, negative metrics'],
        ['Warning Gold', '#C4A86D', 'Caution states'],
    ]
)

add_para('6.2 Typography', bold=True, size=13)
add_table(
    ['Element', 'Font', 'Weight', 'Size'],
    [
        ['Page titles', 'Playfair Display', '700', '30px'],
        ['Section headings', 'Playfair Display', '700', '22px'],
        ['Metric values', 'Outfit', '700', '24-30px'],
        ['Body text', 'Outfit', '400-500', '13-14px'],
        ['Labels', 'Outfit', '600', '10-12px'],
    ]
)

add_para('6.3 Spacing & Effects', bold=True, size=13)
add_bullet('Card border-radius: 14px')
add_bullet('Button border-radius: 10px')
add_bullet('Pill border-radius: 24px')
add_bullet('Card shadow: 0 1px 4px rgba(38,52,49,0.06)')
add_bullet('Elevated shadow: 0 8px 32px rgba(38,52,49,0.10)')

add_para('')
add_para('All tokens live in lib/utils.js as the T object. Every admin page imports T and uses inline CSS-in-JS styles referencing these tokens.', italic=True)

# ============================================================
# 7. ENVIRONMENT VARIABLES
# ============================================================
add_heading('7. Environment Variables', 1)
add_para('Create .env.local with the following keys:')

env_vars = [
    ('Firebase Client (public)', [
        'NEXT_PUBLIC_FIREBASE_API_KEY',
        'NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN',
        'NEXT_PUBLIC_FIREBASE_PROJECT_ID',
        'NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET',
        'NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID',
        'NEXT_PUBLIC_FIREBASE_APP_ID',
    ]),
    ('Firebase Admin (server-only)', [
        'FIREBASE_ADMIN_PROJECT_ID',
        'FIREBASE_ADMIN_CLIENT_EMAIL',
        'FIREBASE_ADMIN_PRIVATE_KEY (include literal \\n)',
    ]),
    ('Razorpay', [
        'NEXT_PUBLIC_RAZORPAY_KEY_ID',
        'RAZORPAY_KEY_ID',
        'RAZORPAY_KEY_SECRET',
    ]),
    ('Domain', [
        'NEXT_PUBLIC_BASE_DOMAIN (default: advertradical.com)',
    ]),
]
for section, vars in env_vars:
    add_para(section, bold=True, size=12)
    for v in vars:
        add_bullet(v)

# ============================================================
# 8. DATABASE FUNCTIONS
# ============================================================
add_heading('8. Key Database Functions (lib/db.js)', 1)
add_para('40+ exported functions organized by domain:')

db_fns = [
    ('Restaurants', 'getRestaurantBySubdomain, getRestaurantById, getAllRestaurants, createRestaurant, updateRestaurant'),
    ('Menu Items', 'getMenuItems, getAllMenuItems, incrementItemView, incrementARView, updateMenuItem, deleteMenuItem, createMenuItem, getAllMenuItemsAllRestaurants'),
    ('Orders', 'createOrder, getOrders, updateOrderStatus, updatePaymentStatus'),
    ('Analytics', 'trackVisit, getAnalytics, getTodayAnalytics, getWaiterCallsCount'),
    ('Waiter Calls', 'createWaiterCall, getWaiterCalls, resolveWaiterCall, deleteWaiterCall'),
    ('Offers', 'getActiveOffers, getAllOffers, createOffer, updateOffer, deleteOffer'),
    ('Combos', 'getCombos, createCombo, updateCombo, deleteCombo'),
    ('Requests', 'getRequests, getAllPendingRequests, submitRequest, updateRequestStatus, deleteRequest, submitRequestAndPublish'),
    ('Users', 'getUserData, createUserDoc'),
    ('Ratings', 'rateMenuItem'),
    ('Table Sessions', 'activateTableSession, clearTableSession, getTableSession, getAllTableSessions, isSessionValid, isSessionValidWithSid'),
]
for domain, fns in db_fns:
    p = doc.add_paragraph()
    run = p.add_run(f'{domain}: ')
    run.bold = True
    p.add_run(fns)

# ============================================================
# 9. ADMIN SIDEBAR
# ============================================================
add_heading('9. Admin Sidebar Navigation', 1)
add_para('The admin sidebar (AdminLayout.jsx) is organized into 5 sections:')

nav_sections = [
    ('OVERVIEW', [('Analytics', '/admin/analytics'), ('Revenue Reports', '/admin/revenue (placeholder)')]),
    ('OPERATIONS', [('Orders', '/admin/orders'), ('Kitchen (KDS)', '/admin/kitchen (placeholder)'), ('Waiter', '/admin/waiter (placeholder)'), ('Payments', '/admin/payments')]),
    ('MENU', [('Menu Items', '/admin/items'), ('Combo Builder', '/admin/combos'), ('Offers', '/admin/offers'), ('Coupons', '/admin/coupons (placeholder)')]),
    ('PEOPLE', [('Staff Logins', '/admin/staff (placeholder)'), ('Customer Feedback', '/admin/feedback (placeholder)'), ('Notifications', '/admin/notifications')]),
    ('SETUP', [('Add Items', '/admin/requests'), ('QR Code', '/admin/qrcode'), ('Settings', '/admin/settings (placeholder)'), ('Subscription', '/admin/subscription')]),
]
for section, items in nav_sections:
    add_para(section, bold=True, size=12)
    for name, route in items:
        add_bullet(f'{name}  →  {route}')

# ============================================================
# 10. ANALYTICS DASHBOARD
# ============================================================
add_heading('10. Analytics Dashboard Details', 1)
add_para('The analytics page (pages/admin/analytics.js, 967 lines) is the most complex admin page. Sections from top to bottom:')

analytics_sections = [
    ('Hero Header', 'Restaurant name + "Analytics" (gold italic), date range selector (7/14/30/90 days)'),
    ('LIVE TODAY', 'Large white card with 4 colored stat cards: Visitors (blue), Orders (green), Revenue (gold), Waiter Calls (coral)'),
    ('Smart Insights', 'Dark green gradient background, 4 AI-generated insights with typed badges (WIN / OPPORTUNITY / ACTION NEEDED / INSIGHT)'),
    ('Customer Journey Funnel', 'Dark background, tapering funnel (Menu Views → AR Engaged → Ordered), gold conversion badges'),
    ('Dish Performance', 'Bento grid: best seller hero card + 5 item cards with progress bars'),
    ('Visits Over Time', 'Gold/green area chart (Recharts AreaChart)'),
    ('Waiter Call Summary', 'Dark cinematic card with response time, resolution rate, total calls'),
    ('Top Menu Items', 'Horizontal bar list showing top 8 items with numbered rankings'),
    ('Restaurant Health Score', 'Score out of 100, motivation quote when >= 70, always-visible alerts or ALL CLEAR'),
]
for i, (name, desc) in enumerate(analytics_sections, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {name}')
    run.bold = True
    p.add_run(f'  —  {desc}')

# ============================================================
# 11. CRITICAL RULES
# ============================================================
add_heading('11. Critical Rules & Conventions', 1)

rules = [
    ('Pages Router ONLY', 'Never use App Router (app/ directory). All routes are in pages/.'),
    ('Complete file outputs', 'When modifying a file, always output the complete file, never partial patches.'),
    ('CSS-in-JS inline styles', 'The project uses inline style={{}} props, not CSS modules or Tailwind classes for component styling.'),
    ('Design tokens', 'Always use T.xxx from lib/utils.js for colors, fonts, radii, shadows. Never hardcode hex values.'),
    ('Dual Firebase', 'Admin and superadmin use separate Firebase app instances. Never cross them.'),
    ('isSoldOutToday guard', 'Menu items have an isSoldOutToday field. Always check it when displaying items to customers.'),
    ('Font hierarchy', 'Page title (30px) > Section heading (22px) > Metric values (24-30px) > Body (13-14px) > Labels (10-12px). Never invert.'),
    ('Typography', 'Playfair Display for headings, Outfit for body and numbers.'),
]
for title, desc in rules:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 12. REMAINING WORK
# ============================================================
add_heading('12. Remaining Work & Roadmap', 1)

add_para('HIGH PRIORITY', bold=True, size=13)
high = [
    'Razorpay payments UI — API routes exist (/api/payments/), subscription page UI needs full wiring',
    'Multi-language menu fields — nameTA (Tamil), nameHI (Hindi) inputs in admin items page',
    'Video in landing page — Chapter 3 section currently shows image, replace with video when asset ready',
    'Placeholder admin pages — Revenue Reports, Kitchen (KDS), Waiter Management, Coupons, Staff Logins, Customer Feedback, Settings',
]
for item in high:
    add_bullet(item)

add_para('')
add_para('MEDIUM PRIORITY', bold=True, size=13)
med = [
    'Allergen/dietary tags — Veg/Non-veg/Jain/Gluten-free badge system',
    'Loyalty & coupons system — Customer loyalty points, coupon generation and redemption',
    'Kitchen Display System (KDS) — Real-time order queue for kitchen staff',
    'Staff management — Multiple staff logins per restaurant with role-based access',
]
for item in med:
    add_bullet(item)

add_para('')
add_para('LOW PRIORITY', bold=True, size=13)
low = [
    'Custom domain — advertradical.com not yet purchased',
    'Dark/light mode toggle — Currently non-functional in admin',
    'Pitch deck visual polish — Awaiting reference designs',
    'PWA support — Service worker for offline menu viewing',
    'Email notifications — Order confirmations, daily summary emails',
]
for item in low:
    add_bullet(item)

add_para('')
add_para('KNOWN ISSUES', bold=True, size=13)
issues = [
    'Some admin sidebar nav items point to pages that do not exist yet',
    'Analytics isSoldOutToday data depends on restaurant actually using the sold-out feature',
    'Table session expiry cleanup is not automated (no cron/cloud function)',
    'No test suite currently exists',
]
for item in issues:
    add_bullet(item)

# ============================================================
# 13. DEPLOYMENT
# ============================================================
add_heading('13. Deployment', 1)
add_para('Vercel auto-deploys from the main branch on every push.')
add_bullet('Build command: next build')
add_bullet('Output directory: .next')
add_bullet('Framework preset: Next.js (auto-detected)')
add_bullet('Domain: ar-saa-s-kbzn.vercel.app')
add_bullet('Custom domain advertradical.com — not yet purchased')
add_para('')
add_para('Local development:', bold=True)
add_code_block('npm install')
add_code_block('# Copy .env.local with Firebase + Razorpay credentials')
add_code_block('npm run dev')
add_code_block('# Visit http://localhost:3000')
add_code_block('# Use ?sub=spot to simulate subdomain routing')

# ============================================================
# 14. THIRD-PARTY SERVICES
# ============================================================
add_heading('14. Third-Party Services', 1)
add_table(
    ['Service', 'Purpose', 'Dashboard'],
    [
        ['Firebase', 'Auth, Firestore, Storage', 'console.firebase.google.com'],
        ['Vercel', 'Hosting & deployment', 'vercel.com/dashboard'],
        ['Razorpay', 'Payment processing', 'dashboard.razorpay.com'],
        ['Meshy AI', '3D model generation', 'meshy.ai'],
        ['Anthropic', 'Claude AI for upselling', 'console.anthropic.com'],
    ]
)

# ============================================================
# 15. GIT WORKFLOW
# ============================================================
add_heading('15. Git Workflow', 1)
add_bullet('Main branch: main — auto-deploys to Vercel')
add_bullet('Working branches: Created per feature/session')
add_bullet('Commit style: Descriptive — "Fix My Bill mobile tap — restructure sheet layout to fix iOS Safari bug"')
add_bullet('No CI/CD pipeline beyond Vercel auto-deploy')
add_bullet('No test suite currently')

# ============================================================
# 16. ONBOARDING
# ============================================================
add_heading('16. How to Onboard a New Developer / AI Assistant', 1)
steps = [
    'Clone the repo and run npm install',
    'Get .env.local credentials from Prabu (Firebase, Razorpay, Meshy keys)',
    'Read this handoff document fully',
    'Read lib/utils.js for design tokens (T object)',
    'Read lib/db.js for all Firestore operations',
    'Read hooks/useAuth.js for auth flow',
    'Check components/layout/AdminLayout.jsx for sidebar structure',
    'Start dev server: npm run dev',
    'Test customer menu at localhost:3000/restaurant/spot',
    'Test admin at localhost:3000/admin/login',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'{i}. {step}')

add_para('')
p = doc.add_paragraph()
run = p.add_run('Key principle: ')
run.bold = True
p.add_run('This project uses CSS-in-JS inline styles with design tokens from the T object. Always maintain the cinematic visual language — Deep Forest, Antique Gold, Soft Cream palette with Playfair Display headings and Outfit body text.')
p.runs[-1].italic = True

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— End of Handoff Document —')
run.italic = True
run.font.color.rgb = RGBColor(0x63, 0x5F, 0x5A)

# Save
output_path = r'C:\Users\Prabu D\OneDrive\Desktop\advert-radical-v6\advert-radical\HANDOFF.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
